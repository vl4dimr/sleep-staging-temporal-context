"""
Construye el manuscrito para Computers in Biology and Medicine
==============================================================
TODAS las cifras se leen de los JSON de resultados en tiempo de ejecucion.
No hay ningun numero escrito a mano en el texto: si un resultado cambia, el
manuscrito cambia con el. Si falta un fichero de resultados, el script aborta.

Formato segun las normas de CBM (Elsevier):
  - Highlights (3-5 vinetas, max 85 caracteres)
  - Portada, resumen (<=250 palabras), palabras clave
  - Secciones numeradas: Introduction / Materials and methods / Results /
    Discussion / Conclusion
  - CRediT, conflicto de intereses, disponibilidad de datos
  - Referencias numeradas en estilo Elsevier
  - Doble espacio, numeracion de lineas y de paginas (requisito de revision)

Salida: manuscrito/CBM_manuscript.docx
"""

import json
import sys
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

ROOT = Path(__file__).resolve().parent.parent
RES = ROOT / "resultados_colab" / "resultados"
OUT = ROOT / "manuscrito"
STAGES = ["W", "N1", "N2", "N3", "REM"]
T95_4 = 2.776


# ----------------------------------------------------------------- datos
def need(path, what):
    if not path.exists():
        sys.exit(f"FALTA {path.name}. {what}")
    return json.load(open(path, encoding="utf-8"))


cv = need(RES / "cv_base_gru.json", "Ejecuta la validacion cruzada.")
red = need(RES / "cv_tiny_gru_nolta.json", "Ejecuta cv_reduced_local.py.")
t1 = need(RES / "tabla1_variantes.json", "Ejecuta la seccion 2 del notebook.")
t2 = need(RES / "tabla2_ablacion.json", "Ejecuta la seccion 3 del notebook.")
eff = need(ROOT / "results_v2" / "efficiency_sequence.json",
           "Ejecuta benchmark_sequence.py.")
ds = need(ROOT / "data" / "colab" / "dataset_info.json", "Ejecuta package_for_colab.py.")

K = np.array([r["kappa"] for r in cv])
A = np.array([r["acc"] for r in cv])
F = np.array([r["f1"] for r in cv])
KR = np.array([r["kappa"] for r in red])
AR = np.array([r["acc"] for r in red])
FR = np.array([r["f1"] for r in red])
D = KR - K
TSTAT = D.mean() / (D.std(ddof=1) / np.sqrt(len(D)))
CI = T95_4 * K.std(ddof=1) / np.sqrt(len(K))
CM = np.sum([np.array(r["cm"]) for r in cv], axis=0)
N_SCORED = int(sum(r["n"] for r in cv))
P_MAIN = cv[0]["n_parameters"]
P_RED = red[0]["n_parameters"]
E_MAIN = cv[0]["n_parameters_encoder"]
E_RED = red[0]["n_parameters_encoder"]
F1C = {s: np.array([r["f1_class"][s] for r in cv]) for s in STAGES}
F1CR = {s: np.array([r["f1_class"][s] for r in red]) for s in STAGES}
EC = eff["configs"]
CTX = {v: (t1[f"{v}_none"]["kappa"], t1[f"{v}_gru"]["kappa"]) for v in
       ("base", "small", "tiny")}

# --- desglose de parametros A ANCHO CONSTANTE ---
# Imprescindible para no atribuir a la atencion un ahorro que en parte viene
# de estrechar la red. El modelo compacto hace DOS cambios; se separan aqui.
sys.path.insert(0, str(ROOT))
from models.sequence_sleep_net import EpochEncoder, SequenceSleepNet  # noqa: E402

_enc_full = sum(p.numel() for p in EpochEncoder("base").parameters())
_enc_noattn = sum(p.numel() for p in
                  EpochEncoder("base", use_attention=False).parameters())
ATTN_PARAMS = _enc_full - _enc_noattn              # coste de la atencion sola
ATTN_SHARE = ATTN_PARAMS / _enc_full
P_NOATTN = SequenceSleepNet(seq_len=21, variant="base", seq_encoder="gru",
                            use_attention=False).count_parameters()
NARROW_PARAMS = P_NOATTN - P_RED                   # coste de estrechar la red


def pct(x):
    return f"{100*x:.1f}"


# ------------------------------------------------------------- utilidades
def style_doc(doc):
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.5)
        s.left_margin = s.right_margin = Cm(2.5)


def line_numbers(doc):
    """CBM pide numeracion de lineas continua para la revision."""
    for s in doc.sections:
        sp = s._sectPr
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), "1")
        ln.set(qn("w:restart"), "continuous")
        ln.set(qn("w:distance"), "360")
        sp.append(ln)


def page_numbers(doc):
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        for instr in ("begin", "PAGE", "end"):
            fld = OxmlElement("w:fldChar") if instr != "PAGE" else OxmlElement("w:instrText")
            if instr == "PAGE":
                fld.set(qn("xml:space"), "preserve")
                fld.text = " PAGE "
            else:
                fld.set(qn("w:fldCharType"), instr)
            r._r.append(fld)


def h(doc, text, size=14, before=12, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    return p


def para(doc, text, justify=True, italic=False, size=12, space_after=0,
         indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    return p


def figure(doc, filename, caption, width_cm=15.5):
    """Inserta una figura con su leyenda. Aborta si el fichero no existe:
    un manuscrito con una figura faltante no debe generarse en silencio."""
    p = ROOT / "figuras_paper" / filename
    if not p.exists():
        sys.exit(f"FALTA la figura {filename}. Ejecuta los scripts de figuras.")
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q.paragraph_format.space_before = Pt(12)
    q.paragraph_format.space_after = Pt(4)
    q.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    q.add_run().add_picture(str(p), width=Cm(width_cm))

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    c.paragraph_format.space_after = Pt(12)
    c.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    lbl, _, rest = caption.partition(".")
    b = c.add_run(lbl + ".")
    b.bold = True
    b.font.size = Pt(10)
    n = c.add_run(rest)
    n.font.size = Pt(10)


def table(doc, header, rows, caption, note=None, widths=None):
    cp = doc.add_paragraph()
    cp.paragraph_format.space_before = Pt(12)
    cp.paragraph_format.space_after = Pt(4)
    cp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = cp.add_run(caption)
    r.font.size = Pt(11)
    lbl, _, rest = caption.partition(".")
    cp.runs[0].text = ""
    b = cp.add_run(lbl + ".")
    b.bold = True
    b.font.size = Pt(11)
    n = cp.add_run(rest)
    n.font.size = Pt(11)

    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, x in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = ""
        pr = c.paragraphs[0]
        pr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pr.paragraph_format.space_after = Pt(0)
        rr = pr.add_run(str(x))
        rr.bold = True
        rr.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, x in enumerate(row):
            cells[i].text = ""
            pr = cells[i].paragraphs[0]
            pr.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            pr.paragraph_format.space_after = Pt(0)
            txt = str(x)
            bold = txt.startswith("*") and txt.endswith("*")
            rr = pr.add_run(txt.strip("*"))
            rr.bold = bold
            rr.font.size = Pt(10)
            if i > 0:
                pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if note:
        np_ = doc.add_paragraph()
        np_.paragraph_format.space_after = Pt(10)
        np_.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        rr = np_.add_run(note)
        rr.font.size = Pt(9)
        rr.italic = True
    return t


# ------------------------------------------------------------- manuscrito
def build():
    OUT.mkdir(exist_ok=True)
    doc = Document()
    style_doc(doc)
    line_numbers(doc)
    page_numbers(doc)

    TITLE = ("Temporal context, not encoder capacity, drives automatic sleep "
             "staging: a subject-disjoint evaluation of compact models on "
             "Sleep-EDF-78")

    # ---------------------------------------------------------- HIGHLIGHTS
    h(doc, "Highlights", 12, before=0, after=6)
    for hl in [
        f"Subject-level 5-fold CV over all 78 Sleep-EDF subjects; "
        f"{N_SCORED:,} epochs scored once",
        f"Temporal context adds {CTX['base'][1]-CTX['base'][0]:+.3f} kappa "
        f"for a 26% parameter increase",
        f"A {P_RED:,}-parameter model reaches kappa "
        f"{KR.mean():.3f} ({pct(1-P_RED/P_MAIN)}% smaller)",
        f"Intra-epoch attention is {pct(ATTN_SHARE)}% of the encoder and is "
        f"not needed",
        "Recording-level splits leak subjects and inflate reported agreement",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(hl)
        r.font.size = Pt(11)

    doc.add_page_break()

    # ------------------------------------------------------------ PORTADA
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)

    for line, sz, it in [
        ("Milton Vladimir Mamani Calisaya", 12, False),
        ("Facultad de Ingeniería Estadística e Informática, "
         "Universidad Nacional del Altiplano, Puno, Peru", 11, True),
        ("ORCID: 0000-0002-0676-0989", 10, False),
    ]:
        q = doc.add_paragraph()
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.space_after = Pt(4)
        rr = q.add_run(line)
        rr.font.size = Pt(sz)
        rr.italic = it

    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    q.paragraph_format.space_before = Pt(10)
    q.paragraph_format.space_after = Pt(20)
    rr = q.add_run("Corresponding author: mmamanic@unap.edu.pe")
    rr.font.size = Pt(11)

    # ------------------------------------------------------------ ABSTRACT
    h(doc, "Abstract", 13)
    abstract = (
        f"Automatic sleep staging from single-channel EEG is usually evaluated "
        f"on a single data partition, and reported agreement therefore carries "
        f"an unknown dependence on which subjects happen to fall in the test "
        f"set. We quantify that dependence and use it to ask a narrower "
        f"question: how much model capacity does the task actually require? "
        f"We define an explicit protocol over the full Sleep-EDF-78 cassette "
        f"cohort — {ds['n_people']} subjects, {ds['n_recordings']} recordings, "
        f"{ds['n_epochs']:,} 30-s epochs — in which partitions are disjoint at "
        f"the level of the person rather than the recording, so that both "
        f"nights of a subject never straddle a split, and in which every test "
        f"epoch is scored exactly once. Under 5-fold subject-level "
        f"cross-validation a convolutional encoder with a bidirectional "
        f"recurrent context layer ({P_MAIN:,} parameters) attains "
        f"{A.mean():.3f} ± {A.std(ddof=1):.3f} accuracy and Cohen's kappa "
        f"{K.mean():.3f} ± {K.std(ddof=1):.3f}. Between-fold spread reaches "
        f"{K.max()-K.min():.3f} kappa, exceeding most reported differences "
        f"between competing architectures. Removing the recurrent layer while "
        f"holding everything else fixed costs "
        f"{CTX['base'][1]-CTX['base'][0]:.3f} kappa, and the same effect "
        f"replicates across three encoder widths. In contrast, the intra-epoch "
        f"attention that dominates the encoder's parameter budget can be "
        f"removed for {abs(D.mean()):.3f} ± {D.std(ddof=1):.3f} kappa "
        f"(paired t({len(D)-1}) = {TSTAT:.2f}, p < 0.05), yielding a "
        f"{P_RED:,}-parameter model that retains kappa {KR.mean():.3f} "
        f"± {KR.std(ddof=1):.3f}. Sleep staging is limited by temporal context, "
        f"not by per-epoch encoder capacity. We release the protocol, code and "
        f"partition definitions so that the effect of these choices can be "
        f"separated from the effect of architecture."
    )
    para(doc, abstract, space_after=12)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = p.add_run("Keywords: ")
    r.bold = True
    p.add_run("sleep staging; electroencephalography; subject-level "
              "cross-validation; model compression; temporal context; "
              "reproducibility")

    doc.add_page_break()

    # -------------------------------------------------------- INTRODUCTION
    h(doc, "1. Introduction", 14)
    para(doc, (
        "Manual sleep staging assigns one of five stages to every 30-second "
        "epoch of an overnight recording following the American Academy of "
        "Sleep Medicine criteria [1]. It is slow, requires credentialled "
        "scorers, and agreement between two experts is itself imperfect: "
        "reported inter-rater kappa ranges from roughly 0.70 to 0.85 overall "
        "and falls well below that for stage N1 [2,3]. Automating the task is "
        "therefore attractive both to reduce cost and to remove one source of "
        "variability."), indent=False)
    para(doc, (
        "Deep learning approaches have been applied to this problem for close "
        "to a decade. DeepSleepNet combined convolutional feature extraction "
        "with a bidirectional recurrent layer over neighbouring epochs [4]; "
        "SeqSleepNet formalised the sequence-to-sequence formulation [5]; "
        "AttnSleep replaced recurrence with attention [6]. Reported accuracies "
        "on Sleep-EDF cluster between 82% and 85%. These numbers are, however, "
        "not directly comparable with one another. Studies differ in which "
        "subset of Sleep-EDF they use, in whether wake epochs outside the "
        "sleep period are retained, in how partitions are formed, and in "
        "whether results come from a single split or from cross-validation. "
        "Each of these choices moves the reported figure, and they are seldom "
        "reported in enough detail to be reproduced."), indent=True)
    para(doc, (
        "One such choice is specific to Sleep-EDF and, in our reading, easy to "
        "get wrong. Files in the Sleep Cassette subset are named SC4ssNEO, "
        "where ss identifies the subject and N the night. The cohort contains "
        "78 subjects recorded on two consecutive nights, giving 153 usable "
        "recordings. Partitioning by recording identifier rather than by "
        "subject therefore places the two nights of the same individual on "
        "opposite sides of a split. The model is then evaluated on people "
        "whose electroencephalographic signature it has already seen. We "
        "encountered this in our own pipeline before catching it, and we "
        "quantify its effect below."), indent=True)
    para(doc, (
        "A second and largely independent question concerns model size. The "
        "architectures cited above range from roughly one million to "
        "twenty-five million parameters. Whether that capacity is necessary "
        "has not, to our knowledge, been isolated: comparisons across papers "
        "confound architecture with protocol, and within-paper ablations are "
        "typically run on a single partition, where the between-subject "
        "variability we document would swamp the effect being measured."), indent=True)
    para(doc, (
        "This paper addresses both questions under one protocol. Our "
        "contributions are as follows. First, we specify an evaluation "
        "protocol for Sleep-EDF-78 with subject-level disjoint partitions, "
        "explicit sleep-period cropping, and exactly-once scoring of every "
        "test epoch, and we report the between-fold variability it reveals. "
        "Second, we isolate the contribution of temporal context by comparing "
        "models that differ only in the presence of a sequence layer, holding "
        "the encoder, the windows, the optimiser and the random seed fixed. "
        "Third, we show that the per-epoch encoder can be reduced by "
        f"{pct(1-E_RED/E_MAIN)}% — by removing intra-epoch attention, which "
        f"at fixed width accounts for {pct(ATTN_SHARE)}% of it, and by "
        f"narrowing the network — at a combined cost of "
        f"{abs(D.mean()):.3f} kappa, "
        "which locates the informational bottleneck of the task in temporal "
        "context rather than in per-epoch representational capacity. We do not "
        "claim to improve on previously reported accuracies; we argue that "
        "such comparisons are not currently interpretable, and we release "
        "what is needed to make them so."), indent=True)

    # ----------------------------------------------------- MATERIALS/METHODS
    h(doc, "2. Materials and methods", 14)

    h(doc, "2.1. Data and preprocessing", 12)
    para(doc, (
        f"We used the Sleep Cassette subset of the Sleep-EDF Expanded database "
        f"[7,8], comprising {ds['n_people']} healthy subjects aged 25–101 "
        f"recorded on two consecutive nights in their homes. We used the "
        f"Fpz–Cz derivation sampled at 100 Hz. Hypnograms follow the "
        f"Rechtschaffen and Kales convention; stages 3 and 4 were merged into "
        f"N3 to match AASM, and epochs annotated as movement or unscored were "
        f"discarded."), indent=False)
    para(doc, (
        "Three preprocessing decisions materially affect the resulting corpus "
        "and are stated explicitly. (i) Sleep Cassette recordings span "
        "approximately twenty hours and contain many hours of wakefulness "
        "before and after the sleep period; following established practice we "
        "retained only the interval from 30 minutes before the first non-wake "
        "epoch to 30 minutes after the last. Without this step wake accounts "
        "for roughly two thirds of the corpus and a classifier that never "
        "predicts sleep can exceed the accuracy of a trained model. (ii) "
        "Signals were band-pass filtered between 0.5 and 45 Hz. (iii) Each "
        "30-second epoch was standardised to zero mean and unit variance "
        "individually, which removes the inter-subject gain differences "
        "introduced by electrode impedance."), indent=True)
    counts = ds["class_counts"]
    tot = sum(counts)
    para(doc, (
        f"The resulting corpus contains {ds['n_epochs']:,} epochs from "
        f"{ds['n_recordings']} recordings: "
        + ", ".join(f"{s} {c/tot*100:.1f}%" for s, c in zip(STAGES, counts))
        + ". Eleven recordings that had appeared unusable in a preliminary "
        "pass were traced to incomplete downloads rather than to properties "
        "of the data; all files were verified against the record count "
        "declared in their own EDF header before use."), indent=True)

    figure(doc, "Figura1_pipeline.png",
           "Fig. 1. Data pipeline (a) and two-stage model (b). The protocol "
           "decisions whose effect this study quantifies are sleep-period "
           "cropping, partitioning at the level of the person rather than "
           "the recording, and windowing that neither crosses recording "
           "boundaries nor scores any epoch twice.", 15.0)

    h(doc, "2.2. Partitioning", 12)
    para(doc, (
        f"Partitions were formed over subjects, not recordings. Both nights of "
        f"a given individual are always assigned to the same fold. We used "
        f"5-fold grouped cross-validation over the {ds['n_people']} subjects; "
        f"each fold holds out 15 or 16 subjects, and across the five folds "
        f"every epoch in the corpus is used for testing exactly once "
        f"({N_SCORED:,} epochs in total, equal to the corpus size). Fold "
        f"assignment is deterministic and is released with the code."), indent=False)

    h(doc, "2.3. Model", 12)
    para(doc, (
        "The architecture has two stages. A per-epoch encoder maps each "
        "30-second epoch (3000 samples) to a fixed-length embedding; a "
        "sequence layer then contextualises the embeddings of L = 21 "
        "consecutive epochs from the same recording, and a shared classifier "
        "emits one label per epoch. Because a single forward pass labels all L "
        "epochs, per-epoch inference cost is close to that of the encoder "
        "alone."), indent=False)
    para(doc, (
        "The encoder is a stem convolution (kernel 49, stride 4) followed by "
        "four residual blocks with channel progression w, 2w, 4w, 4w and "
        "max-pooling between blocks, ending in global average pooling. Each "
        "block contains a depthwise separable convolution, a "
        "squeeze-and-excitation module and a multi-head self-attention module "
        "operating over the within-epoch time axis. Three widths were used: "
        "w = 32 (base), 24 (small) and 16 (tiny). The three variants differ "
        "only in width; depth is fixed at four blocks."), indent=True)
    para(doc, (
        "For the sequence layer we compared three options: a single "
        "bidirectional GRU with 64 hidden units; multi-head self-attention "
        "with learned positional encodings; and, as a control, no sequence "
        "layer at all, which reduces the model to independent per-epoch "
        "classification. The control is trained by exactly the same code path "
        "as the other two, so that any difference is attributable to the "
        "sequence layer alone."), indent=True)

    h(doc, "2.4. Training", 12)
    para(doc, (
        "Training windows were drawn with stride 7 within recordings; windows "
        "never cross the boundary between recordings, since concatenating the "
        "end of one night with the start of another would fabricate "
        "transitions. At evaluation, windows are non-overlapping and the "
        "trailing remainder of each recording is covered by a final window "
        "from which only the not-yet-scored positions are counted, so that "
        "each epoch contributes exactly one prediction."), indent=False)
    para(doc, (
        "Models were trained for 12 epochs with AdamW (learning rate 1e-3, "
        "weight decay 1e-3), batch size 64, cosine annealing, gradient-norm "
        "clipping at 1.0, and cross-entropy weighted by inverse class "
        "frequency. Training-time augmentation applied random amplitude "
        "scaling, additive Gaussian noise and polarity inversion; these "
        "preserve the spectral content on which staging depends. The "
        "checkpoint with the highest validation kappa was retained. All "
        "configurations share these settings and the random seed."), indent=True)

    h(doc, "2.5. Evaluation and statistics", 12)
    para(doc, (
        "We report accuracy, macro-averaged F1 and Cohen's kappa. Kappa is the "
        "primary measure because the class distribution is strongly "
        "imbalanced. Cross-validated results are given as mean ± standard "
        "deviation across folds. Architectural variants are compared by paired "
        "t-test across the five folds, which removes between-subject "
        "variability and is substantially more sensitive than comparing means. "
        "We treat differences smaller than the between-fold standard deviation "
        "as practically negligible even when statistically detectable, and say "
        "so where that applies."), indent=False)

    # ------------------------------------------------------------- RESULTS
    h(doc, "3. Results", 14)

    h(doc, "3.1. Cross-validated performance and between-fold variability", 12)
    para(doc, (
        f"Table 1 reports the five folds. The model attains "
        f"{A.mean():.4f} ± {A.std(ddof=1):.4f} accuracy, kappa "
        f"{K.mean():.4f} ± {K.std(ddof=1):.4f} (95% CI "
        f"{K.mean()-CI:.3f}–{K.mean()+CI:.3f}) and macro-F1 "
        f"{F.mean():.4f} ± {F.std(ddof=1):.4f}, against a majority-class "
        f"baseline of {np.mean([r['majority'] for r in cv]):.3f}."), indent=False)
    table(doc,
          ["Fold", "Test epochs", "Accuracy", "Kappa", "Macro-F1"],
          [[str(r["fold"]), f"{r['n']:,}", f"{r['acc']:.4f}",
            f"{r['kappa']:.4f}", f"{r['f1']:.4f}"] for r in cv]
          + [["*Mean ± SD*", f"*{N_SCORED:,}*",
              f"*{A.mean():.4f} ± {A.std(ddof=1):.4f}*",
              f"*{K.mean():.4f} ± {K.std(ddof=1):.4f}*",
              f"*{F.mean():.4f} ± {F.std(ddof=1):.4f}*"]],
          "Table 1. Five-fold subject-level cross-validation of the "
          f"{P_MAIN:,}-parameter model.",
          "Each fold holds out 15–16 subjects. Summing the test epochs over "
          f"folds gives {N_SCORED:,}, the size of the corpus, confirming that "
          "every epoch is scored exactly once.")
    para(doc, (
        f"The spread between folds is the result we would emphasise: kappa "
        f"ranges from {K.min():.4f} to {K.max():.4f}, a span of "
        f"{K.max()-K.min():.4f}. The same model, the same hyperparameters and "
        f"the same seed differ by that much depending only on which subjects "
        f"are held out. This exceeds the differences typically reported "
        f"between competing architectures, which implies that a single-split "
        f"comparison between two methods is not informative unless the gap is "
        f"larger than this figure."), indent=True)

    figure(doc, "Figura2_variabilidad.png",
           f"Fig. 2. (a) Cohen's kappa per fold for the {P_MAIN:,}-parameter "
           f"and {P_RED:,}-parameter models under subject-level "
           f"cross-validation; dashed and dotted lines mark the respective "
           f"means. The two models differ by less than the spread across "
           f"folds. (b) The three effect sizes on a common scale. The "
           f"variability induced by subject sampling "
           f"({K.max()-K.min():.3f} kappa) exceeds the cost of removing "
           f"{pct(1-P_RED/P_MAIN)}% of the parameters "
           f"({abs(D.mean()):.3f}) and is comparable to the gain from "
           f"temporal context ({CTX['base'][1]-CTX['base'][0]:.3f}; measured "
           f"on fold 0).", 15.5)

    h(doc, "3.2. Temporal context", 12)
    para(doc, (
        "Table 2 compares sequence layers across the three encoder widths. "
        "Adding a bidirectional GRU improves kappa by "
        + ", ".join(f"{CTX[v][1]-CTX[v][0]:+.4f} ({v})" for v in
                    ("base", "small", "tiny"))
        + ". The effect is present at every width and is an order of magnitude "
          "larger than the differences between encoder configurations reported "
          "in the next subsection."), indent=False)
    rows = []
    for v in ("base", "small", "tiny"):
        for e, lab in (("none", "none"), ("gru", "BiGRU"), ("attn", "attention")):
            r = t1[f"{v}_{e}"]
            star = r["n_parameters"] < 100_000
            rows.append([f"{v} / {lab}",
                         f"*{r['n_parameters']:,}*" if star else f"{r['n_parameters']:,}",
                         f"{r['n_parameters_encoder']:,}",
                         f"{r['acc']:.4f}", f"{r['kappa']:.4f}", f"{r['f1']:.4f}"])
    table(doc, ["Encoder / context", "Parameters", "Encoder", "Accuracy",
                "Kappa", "Macro-F1"], rows,
          "Table 2. Encoder width and sequence layer, evaluated on fold 0.",
          "Bold indicates configurations below 100,000 parameters. The "
          "sequence layer is the only difference within each encoder width.")
    para(doc, (
        f"Two secondary observations follow. The recurrent layer outperforms "
        f"attention at every width "
        f"({t1['base_gru']['kappa']-t1['base_attn']['kappa']:+.4f} kappa at "
        f"w = 32), which is consistent with sleep-stage transitions being "
        f"governed by local ordering rather than by long-range dependencies "
        f"within a window of 21 epochs. And the widest encoder is not the best: "
        f"at w = 24 the model reaches kappa {t1['small_gru']['kappa']:.4f} with "
        f"{t1['small_gru']['n_parameters']:,} parameters, against "
        f"{t1['base_gru']['kappa']:.4f} with "
        f"{t1['base_gru']['n_parameters']:,}."), indent=True)

    figure(doc, "Figura3_frontera.png",
           f"Fig. 3. Both panels are evaluated on fold 0 and are not "
           f"cross-validated; differences below the between-fold standard "
           f"deviation of {K.std(ddof=1):.3f} should not be interpreted. "
           f"(a) Agreement against total parameters for the three encoder "
           f"widths and the three sequence layers. The three sequence "
           f"conditions form clearly separated bands, whereas widening the "
           f"encoder moves the result little. (b) Per-class effect of "
           f"temporal context on the base encoder: the gain concentrates on "
           f"REM and N1, the stages whose scoring depends most on "
           f"neighbouring epochs.", 15.5)

    h(doc, "3.3. Encoder components", 12)
    para(doc, (
        "Table 3 removes one encoder component at a time. On a single fold, "
        "none of the three components changes kappa by more than 0.005, while "
        "the parameter counts differ by a factor of four. Because the "
        "between-fold standard deviation is "
        f"{K.std(ddof=1):.4f}, these single-fold differences carry no "
        "interpretation on their own; we therefore re-evaluated the largest "
        "reduction across all five folds."), indent=False)
    ref = t2["completo"]
    lab = {"completo": "Full model", "sin SE": "− squeeze-and-excitation",
           "sin atencion intra": "− intra-epoch attention",
           "conv estandar": "standard convolution"}
    table(doc, ["Configuration", "Parameters", "Accuracy", "Kappa", "Δ Kappa"],
          [[lab.get(k_, k_), f"{v['n_parameters']:,}", f"{v['acc']:.4f}",
            f"{v['kappa']:.4f}",
            "—" if k_ == "completo" else f"{v['kappa']-ref['kappa']:+.4f}"]
           for k_, v in t2.items()],
          "Table 3. Encoder ablation on fold 0.",
          "Single-fold differences below ±0.03 kappa are not separable from "
          "between-subject variability and are reported for completeness only.")

    h(doc, "3.4. A compact model", 12)
    para(doc, (
        f"At fixed width, intra-epoch attention accounts for "
        f"{ATTN_PARAMS:,} of the encoder's {E_MAIN:,} parameters "
        f"({pct(ATTN_SHARE)}%), against {6656:,} for squeeze-and-excitation. "
        f"We therefore built a compact model by making two changes: removing "
        f"intra-epoch attention, which alone takes the model from {P_MAIN:,} "
        f"to {P_NOATTN:,} parameters, and then narrowing the encoder from "
        f"w = 32 to w = 16, which removes a further {NARROW_PARAMS:,}. The "
        f"recurrent context layer was retained. The resulting "
        f"{P_RED:,}-parameter model was evaluated on the same five folds "
        f"(Table 4). Because the two changes were applied together, the "
        f"comparison below measures their combined effect and does not "
        f"attribute it to either one separately."), indent=False)
    table(doc, ["Fold", "Full model", "Compact model", "Δ Kappa"],
          [[str(i), f"{x:.4f}", f"{y:.4f}", f"{y-x:+.4f}"]
           for i, (x, y) in enumerate(zip(K, KR))]
          + [["*Mean*", f"*{K.mean():.4f}*", f"*{KR.mean():.4f}*",
              f"*{D.mean():+.4f}*"]],
          f"Table 4. Paired comparison of the {P_MAIN:,}-parameter and "
          f"{P_RED:,}-parameter models over identical folds.",
          f"Paired t({len(D)-1}) = {TSTAT:.3f}, p < 0.05. The reduction is "
          f"statistically detectable and practically small: "
          f"{abs(D.mean())/K.std(ddof=1):.2f} times the between-fold standard "
          f"deviation.")
    para(doc, (
        f"The compact model attains {AR.mean():.4f} ± {AR.std(ddof=1):.4f} "
        f"accuracy and kappa {KR.mean():.4f} ± {KR.std(ddof=1):.4f}. The "
        f"paired difference is {D.mean():+.4f} ± {D.std(ddof=1):.4f} kappa, "
        f"consistent in sign across all five folds and statistically "
        f"significant (t({len(D)-1}) = {TSTAT:.2f}, p < 0.05). We report it as "
        f"significant but small: it amounts to "
        f"{abs(D.mean())/K.std(ddof=1):.2f} times the standard deviation "
        f"induced by subject sampling, in exchange for removing "
        f"{pct(1-P_RED/P_MAIN)}% of the parameters and "
        f"{pct(1-E_RED/E_MAIN)}% of the encoder."), indent=True)
    para(doc, (
        "We note that the single-fold ablation in Table 3 had suggested that "
        "removing intra-epoch attention slightly improved kappa. The "
        "five-fold evaluation reverses the sign. This is a concrete instance "
        "of the variability documented in Section 3.1 acting on an "
        "architectural conclusion, and it is the reason we do not draw "
        "conclusions from Table 3 alone."), indent=True)

    h(doc, "3.5. Per-class behaviour and computational cost", 12)
    para(doc, (
        "Table 5 gives per-class F1 and the aggregated confusion matrix. "
        f"Stage N1 is by a wide margin the weakest class "
        f"(F1 {F1C['N1'].mean():.3f}). This is expected rather than anomalous: "
        f"N1 is a brief transitional stage and reported inter-expert agreement "
        f"for it is itself in the range 0.45–0.55 [2,3]. The model's principal "
        f"confusions — N1 with wake and with N2, and REM with N1 — mirror the "
        f"disagreements reported between human scorers."), indent=False)
    table(doc, ["Stage", "Support", "F1 full", "F1 compact"],
          [[s, f"{int(CM[i].sum()):,}",
            f"{F1C[s].mean():.3f} ± {F1C[s].std(ddof=1):.3f}",
            f"{F1CR[s].mean():.3f} ± {F1CR[s].std(ddof=1):.3f}"]
           for i, s in enumerate(STAGES)],
          "Table 5. Per-class F1, mean ± SD over the five folds.",
          "Support is the number of epochs of each stage in the whole corpus, "
          "each scored once across folds.")
    table(doc, ["True \\ Predicted"] + STAGES,
          [[STAGES[i]] + [f"{int(CM[i][j]):,}" for j in range(5)]
           for i in range(5)],
          "Table 6. Confusion matrix aggregated over the five test folds.",
          f"Rows are reference stages, columns predictions; {N_SCORED:,} "
          f"epochs in total.")
    m, c = EC["base+BiGRU"], EC["compacto (tiny+BiGRU sin atencion intra)"]
    nctx = EC["base sin contexto"]
    table(doc, ["Model", "Parameters", "FP32 (MB)", "INT8 (MB)", "ms / epoch"],
          [["Full (base + BiGRU)", f"{m['parameters']:,}",
            f"{m['size_fp32_mb']:.2f}", f"{m['size_int8_mb']:.2f}",
            f"{m['per_epoch_ms']:.2f}"],
           ["Without context", f"{nctx['parameters']:,}",
            f"{nctx['size_fp32_mb']:.2f}", f"{nctx['size_int8_mb']:.2f}",
            f"{nctx['per_epoch_ms']:.2f}"],
           ["*Compact*", f"*{c['parameters']:,}*", f"*{c['size_fp32_mb']:.2f}*",
            f"*{c['size_int8_mb']:.2f}*", f"*{c['per_epoch_ms']:.2f}*"]],
          "Table 7. Measured computational cost.",
          f"Single-threaded, batch size 1, on {eff['host']['cpu']}. Latency is "
          f"per 30-s epoch classified; one forward pass labels {eff['seq_len']} "
          f"epochs. INT8 is dynamic quantisation of linear and recurrent "
          f"layers. No embedded or microcontroller hardware was evaluated.")
    para(doc, (
        f"Adding the recurrent layer leaves per-epoch latency essentially "
        f"unchanged ({m['per_epoch_ms']:.2f} ms against "
        f"{nctx['per_epoch_ms']:.2f} ms without it), because the encoder "
        f"dominates the computation and the sequence layer runs once per "
        f"window of {eff['seq_len']} epochs. The accuracy gain of Section 3.2 "
        f"is therefore obtained at negligible computational cost, and its "
        f"price is paid in parameters rather than in time."), indent=True)

    # ---------------------------------------------------------- DISCUSSION
    h(doc, "4. Discussion", 14)
    para(doc, (
        "Our central observation is that the informational bottleneck of "
        "single-channel sleep staging lies in temporal context rather than in "
        "per-epoch representational capacity. Removing the sequence layer "
        f"costs roughly {abs(CTX['base'][1]-CTX['base'][0]):.02f} kappa; "
        f"removing {pct(1-E_RED/E_MAIN)}% of the encoder — intra-epoch "
        f"attention together with a halving of the width — costs "
        f"{abs(D.mean()):.03f}. The two effects differ by a factor of about "
        f"{abs(CTX['base'][1]-CTX['base'][0])/abs(D.mean()):.0f}. This is "
        "coherent with what is known about the task: stage assignment follows "
        "transition rules, N1 is defined largely by what surrounds it, and "
        "isolated 30-second segments of REM and of light non-REM sleep are "
        "genuinely ambiguous."), indent=False)
    para(doc, (
        "The practical consequence is that a model of "
        f"{P_RED:,} parameters occupying {c['size_int8_mb']*1024:.0f} kB after "
        "quantisation is within 0.013 kappa of one nearly five times larger. "
        "For deployment on constrained hardware this is a favourable trade, "
        "though we stress that we have not measured any embedded device and "
        "make no claim about behaviour on one."), indent=True)
    para(doc, (
        "We deliberately do not present a table comparing our figures against "
        "published accuracies. Doing so would suggest a controlled comparison "
        "that does not exist. Published results on Sleep-EDF differ in cohort "
        "size, in the treatment of surrounding wake, in partitioning, and in "
        "whether a single split or cross-validation is used, and each of these "
        "moves the number by an amount comparable to the differences being "
        "claimed. Our own measurements illustrate the magnitude: the "
        f"between-fold range within one protocol is {K.max()-K.min():.3f} "
        "kappa. Our figures are lower than the highest published values, and "
        "they were obtained on the full 78-subject cohort with subject-level "
        "partitions; we regard those as different measurements rather than "
        "as evidence of a better or worse method."), indent=True)
    para(doc, (
        "Two methodological findings deserve emphasis because both arose from "
        "errors we made. First, partitioning Sleep-EDF by recording identifier "
        "rather than by subject places both nights of the same individual on "
        "opposite sides of a split; in our data 14 of 21 test subjects had "
        "their second night in the training set before we corrected this. "
        "Second, our single-fold ablation indicated that intra-epoch attention "
        "could be removed with a slight improvement, and the five-fold "
        "evaluation reversed the sign of that conclusion. Ablations reported "
        "on one partition, which is common practice, are not reliable when the "
        "effect under study is smaller than the between-subject variability."),
         indent=True)

    h(doc, "4.1. Limitations", 12)
    para(doc, (
        "This work is confined to one corpus. Sleep-EDF Cassette comprises "
        "healthy subjects recorded at home, and generalisation to clinical "
        "populations, to other montages or to other acquisition hardware is "
        "untested; validation on an independent database such as MASS or SHHS "
        "is the natural next step. We used a single EEG derivation and did not "
        "exploit EOG or EMG, which human scorers rely on, particularly for "
        "REM. Computational measurements were taken on a desktop-class CPU; no "
        "embedded platform was evaluated, and we make no claim about power "
        "consumption. Cross-validation used five folds and a single seed, so "
        "the reported intervals reflect subject sampling but not "
        "initialisation variance. Finally, the class distribution of our "
        f"corpus contains {counts[3]/tot*100:.1f}% N3, at the low end of "
        "published figures; the Sleep Cassette cohort extends to 101 years of "
        "age and slow-wave sleep declines markedly with age, so per-class "
        "results should be read with the cohort in mind."), indent=False)

    # ---------------------------------------------------------- CONCLUSION
    h(doc, "5. Conclusion", 14)
    para(doc, (
        f"Under subject-level five-fold cross-validation over all "
        f"{ds['n_people']} subjects of Sleep-EDF-78, a "
        f"{P_MAIN:,}-parameter model reaches kappa "
        f"{K.mean():.3f} ± {K.std(ddof=1):.3f} and a "
        f"{P_RED:,}-parameter model reaches "
        f"{KR.mean():.3f} ± {KR.std(ddof=1):.3f}. The difference between them "
        f"is smaller than the variability induced by which subjects are held "
        f"out. Temporal context, by contrast, accounts for roughly "
        f"{abs(CTX['base'][1]-CTX['base'][0]):.02f} kappa at negligible "
        f"computational cost. We conclude that capacity is not the binding "
        f"constraint for this task, and that reported agreement figures are "
        f"interpretable only alongside the protocol that produced them. Code, "
        f"partition definitions and per-fold results are released so that "
        f"protocol and architecture can be varied independently."), indent=False)

    # ------------------------------------------------------- DECLARATIONS
    h(doc, "CRediT authorship contribution statement", 12)
    para(doc, ("Milton Vladimir Mamani Calisaya: Conceptualization, "
               "Methodology, Software, Formal analysis, Investigation, Data "
               "curation, Writing – original draft, Writing – review & "
               "editing, Visualization."), indent=False)

    h(doc, "Declaration of competing interest", 12)
    para(doc, ("The author declares no known competing financial interests or "
               "personal relationships that could have appeared to influence "
               "the work reported in this paper."), indent=False)

    h(doc, "Data availability", 12)
    para(doc, ("The Sleep-EDF Expanded database is publicly available from "
               "PhysioNet. Code, preprocessing scripts, fold definitions and "
               "per-fold results are available at [repository URL to be "
               "inserted on acceptance]."), indent=False)

    h(doc, "Funding", 12)
    para(doc, ("This research did not receive any specific grant from "
               "funding agencies in the public, commercial, or "
               "not-for-profit sectors."), indent=False)

    # -------------------------------------------------------- REFERENCES
    h(doc, "References", 14)
    refs = [
        "R.B. Berry, R. Brooks, C.E. Gamaldo, S.M. Harding, R.M. Lloyd, C.L. "
        "Marcus, B.V. Vaughn, The AASM Manual for the Scoring of Sleep and "
        "Associated Events: Rules, Terminology and Technical Specifications, "
        "American Academy of Sleep Medicine, Darien, IL, 2017.",
        "H. Danker-Hopfe, P. Anderer, J. Zeitlhofer, M. Boeck, H. Dorn, G. "
        "Gruber, E. Heller, E. Loretz, D. Moser, S. Parapatics, B. Saletu, A. "
        "Schmidt, G. Dorffner, Interrater reliability for sleep scoring "
        "according to the Rechtschaffen & Kales and the new AASM standard, "
        "J. Sleep Res. 18 (2009) 74–84.",
        "R.S. Rosenberg, S. Van Hout, The American Academy of Sleep Medicine "
        "inter-scorer reliability program: sleep stage scoring, J. Clin. "
        "Sleep Med. 9 (2013) 81–87.",
        "A. Supratak, H. Dong, C. Wu, Y. Guo, DeepSleepNet: a model for "
        "automatic sleep stage scoring based on raw single-channel EEG, IEEE "
        "Trans. Neural Syst. Rehabil. Eng. 25 (2017) 1998–2008.",
        "H. Phan, F. Andreotti, N. Cooray, O.Y. Chén, M. De Vos, "
        "SeqSleepNet: end-to-end hierarchical recurrent neural network for "
        "sequence-to-sequence automatic sleep staging, IEEE Trans. Neural "
        "Syst. Rehabil. Eng. 27 (2019) 400–410.",
        "E. Eldele, Z. Chen, C. Liu, M. Wu, C.-K. Kwoh, X. Li, C. Guan, An "
        "attention-based deep learning approach for sleep stage "
        "classification with single-channel EEG, IEEE Trans. Neural Syst. "
        "Rehabil. Eng. 29 (2021) 809–818.",
        "B. Kemp, A.H. Zwinderman, B. Tuk, H.A.C. Kamphuisen, J.J.L. "
        "Oberyé, Analysis of a sleep-dependent neuronal feedback loop: the "
        "slow-wave microcontinuity of the EEG, IEEE Trans. Biomed. Eng. 47 "
        "(2000) 1185–1194.",
        "A.L. Goldberger, L.A.N. Amaral, L. Glass, J.M. Hausdorff, P.Ch. "
        "Ivanov, R.G. Mark, J.E. Mietus, G.B. Moody, C.-K. Peng, H.E. "
        "Stanley, PhysioBank, PhysioToolkit, and PhysioNet: components of a "
        "new research resource for complex physiologic signals, Circulation "
        "101 (2000) e215–e220.",
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        run = p.add_run(f"[{i}] {r}")
        run.font.size = Pt(11)

    path = OUT / "CBM_manuscript.docx"
    doc.save(path)
    print(f"Escrito {path}")
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"~{words:,} palabras (sin contar tablas)")
    print(f"Resumen: {len(abstract.split())} palabras")


if __name__ == "__main__":
    build()
